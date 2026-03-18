# -*- coding: utf-8 -*-
"""
将源程序按模块整理到 Word 文档中
要求：各模块前、后各连续30页（每页≥50行，不含空行），不足3000行全部提交
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 项目根目录
PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOC_PATH = PROJECT_ROOT / "docs" / "作物灾害预警系统使用说明.docx"
# 输出到原文档（脚本会先删除已有的源程序清单部分，再添加，不会重复）
OUTPUT_PATH = PROJECT_ROOT / "docs" / "作物灾害预警系统使用说明.docx"

LINES_PER_PAGE = 50
PAGES_FRONT = 30
PAGES_BACK = 30
LINES_FRONT = PAGES_FRONT * LINES_PER_PAGE  # 1500
LINES_BACK = PAGES_BACK * LINES_PER_PAGE    # 1500
TOTAL_THRESHOLD = 3000  # 不足3000行全部提交


def get_non_empty_lines(content: str) -> list[str]:
    """获取非空行列表"""
    return [line for line in content.splitlines() if line.strip()]


def read_file_lines(filepath: Path) -> tuple[list[str], list[str]]:
    """读取文件，返回 (所有行, 非空行)"""
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        all_lines = f.readlines()
    non_empty = [line for line in all_lines if line.strip()]
    return all_lines, non_empty


def extract_code_for_module(files_data: list[tuple[Path, str]], total_non_empty: int) -> str:
    """
    根据规则提取代码：不足3000行全部提交，否则前30页+后30页
    最后一页应是某段程序的结束页（尽量在文件边界结束）
    """
    if total_non_empty < TOTAL_THRESHOLD:
        # 全部提交
        result = []
        for filepath, content in files_data:
            result.append(f"\n// ========== {filepath} ==========\n")
            result.append(content)
        return "".join(result)

    # 前30页 (1500行) + 后30页 (1500行)
    lines_collected = []
    for filepath, content in files_data:
        non_empty = get_non_empty_lines(content)
        for line in non_empty:
            lines_collected.append((filepath, line))

    # 前1500行
    front_lines = []
    for i, (fp, line) in enumerate(lines_collected[:LINES_FRONT]):
        if i == 0 or (i > 0 and lines_collected[i-1][0] != fp):
            front_lines.append(f"\n// ========== {fp} ==========\n")
        front_lines.append(line + "\n" if not line.endswith("\n") else line)

    # 后1500行 - 从 total - 1500 开始
    start_back = max(LINES_FRONT, len(lines_collected) - LINES_BACK)
    back_lines = []
    for i in range(start_back, len(lines_collected)):
        fp, line = lines_collected[i]
        if i == start_back or lines_collected[i-1][0] != fp:
            back_lines.append(f"\n// ========== {fp} ==========\n")
        back_lines.append(line + "\n" if not line.endswith("\n") else line)

    result = "".join(front_lines)
    result += "\n\n// ... (中间部分省略) ...\n\n"
    result += "".join(back_lines)
    return result


def add_code_to_document(doc: Document, module_name: str, code_content: str):
    """将代码块添加到 Word 文档"""
    # 添加模块标题
    h = doc.add_heading(module_name, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加代码段落 - 使用等宽字体
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_content)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 中文字体
    run.font.name = "Consolas"


# 定义模块：按网页功能细化分类
# 1.相关数据 2.灾害实时监测 3.智能分析 4.灾害预警 5.智慧决策 6.数据可视化 7.用户权限管理
MODULES = [
    ("模块一：入口与配置", [
        "src/main.ts",
        "src/App.vue",
        "vite.config.ts",
        "src/router/index.ts",
        "src/utils/http.ts",
    ]),
    ("模块二：相关数据模块", [
        "src/views/user/RelatedData.vue",
    ]),
    ("模块三：灾害实时监测模块", [
        "src/views/user/MapVisualization.vue",
    ]),
    ("模块四：智能分析模块", [
        "src/views/user/DataAnalysis.vue",
        "src/api/analysis.ts",
        "server/app.py",
    ]),
    ("模块五：灾害预警模块", [
        "src/views/user/WarningSystem.vue",
        "src/stores/data.ts",
    ]),
    ("模块六：智慧决策模块", [
        "src/views/user/DecisionSupport.vue",
    ]),
    ("模块七：数据可视化模块", [
        "src/views/user/Home.vue",
    ]),
    ("模块八：用户权限管理", [
        "src/views/user/Login.vue",
        "src/stores/user.ts",
        "src/composables/useGlobalSearch.ts",
        "src/layouts/AppLayout.vue",
        "src/views/user/About.vue",
    ]),
    ("模块九：Mock 与后端服务", [
        "src/mock/server.ts",
    ]),
]


def _get_element_text(elem) -> str:
    """从 OXML 元素中提取文本"""
    texts = []
    for t in elem.iter():
        tag = t.tag.split("}")[-1] if "}" in str(t.tag) else str(t.tag)
        if tag == "t":
            if t.text:
                texts.append(t.text)
    return "".join(texts)


def remove_existing_source_section(doc: Document) -> None:
    """删除文档中已有的「源程序清单」及之后的所有内容，避免重复添加"""
    body = doc.element.body
    to_remove = []
    found = False
    for child in body:
        if found:
            to_remove.append(child)
            continue
        text = _get_element_text(child)
        if "源程序清单" in text:
            found = True
            to_remove.append(child)
    for elem in to_remove:
        body.remove(elem)


def main():
    if not DOC_PATH.exists():
        print(f"错误：找不到文档 {DOC_PATH}")
        return

    doc = Document(str(DOC_PATH))

    # 先删除已有的源程序清单部分，避免重复
    remove_existing_source_section(doc)

    # 添加分页和源程序总览标题
    doc.add_page_break()
    doc.add_heading("源程序清单", level=1)
    doc.add_paragraph(
        "以下按网页功能细化分类提交源程序："
        "1.相关数据 2.灾害实时监测 3.智能分析 4.灾害预警 5.智慧决策 6.数据可视化 7.用户权限管理。"
        "每页程序≥50行（不含空行）；不足3000行的模块全部提交，超过3000行的模块提交前30页与后30页。"
    )

    for module_name, file_paths in MODULES:
        files_data = []
        total_non_empty = 0

        for rel_path in file_paths:
            full_path = PROJECT_ROOT / rel_path
            if not full_path.exists():
                print(f"跳过不存在的文件: {rel_path}")
                continue
            with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            non_empty_count = len(get_non_empty_lines(content))
            total_non_empty += non_empty_count
            files_data.append((rel_path, content))

        if not files_data:
            continue

        # 提取代码
        code = extract_code_for_module(files_data, total_non_empty)

        # 分块添加（每块约50行，便于排版）
        lines = code.split("\n")
        chunk_size = LINES_PER_PAGE
        chunks = [lines[i:i + chunk_size] for i in range(0, len(lines), chunk_size)]

        doc.add_page_break()
        h = doc.add_heading(module_name, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"本模块共 {total_non_empty} 行有效代码（不含空行）")

        for chunk in chunks:
            chunk_text = "\n".join(chunk)
            if not chunk_text.strip():
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(chunk_text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            r = run._element
            if r.rPr is None:
                r.get_or_add_rPr()
            try:
                r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            except Exception:
                pass

    try:
        doc.save(str(OUTPUT_PATH))
        print(f"已保存到: {OUTPUT_PATH}")
    except PermissionError:
        alt_path = PROJECT_ROOT / "docs" / "作物灾害预警系统使用说明_含源码.docx"
        doc.save(str(alt_path))
        print(f"原文件被占用（可能正在 Word 中打开），已保存到: {alt_path}")
        print("请关闭 Word 中的文档后重试，或直接使用新生成的文件。")


if __name__ == "__main__":
    main()
